from paddleocr import PPStructureV3
from docx import Document
from docx.shared import Pt, Inches
from bs4 import BeautifulSoup
from docx.enum.text import WD_COLOR_INDEX
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
import io

def add_html_table_to_doc(doc, html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    rows = soup.find_all('tr')
    if not rows:
        return
    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < max_cols:
                table.cell(i, j).text = cell.get_text(strip=True)

def render_formula_to_image(doc, latex_content):
    cleaned = latex_content.replace('\\begin{aligned}', '').replace('\\end{aligned}', '')
    lines = cleaned.split('\\\\')

    for line in lines:
        line = line.replace('&', ' ').replace('\\quad', '   ').strip()
        if not line:
            continue

        fig = plt.figure(figsize=(6, 0.6))
        fig.text(0.01, 0.5, f'${line}$', fontsize=16, va='center')
        plt.axis('off')

        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight', dpi=150, transparent=True)
        plt.close(fig)
        buf.seek(0)

        doc.add_picture(buf, width=Inches(4))

def boxes_overlap(box1, box2):
    """Check if two bounding boxes overlap. box format: [x1, y1, x2, y2]"""
    x1_min, y1_min, x1_max, y1_max = box1
    x2_min, y2_min, x2_max, y2_max = box2
    return not (x1_max < x2_min or x2_max < x1_min or y1_max < y2_min or y2_max < y1_min)

def get_min_confidence_for_block(block_bbox, ocr_lines):
    """Find the lowest confidence score among OCR lines that overlap this block's bbox."""
    scores_in_block = []
    for line_box, score in ocr_lines:
        if boxes_overlap(block_bbox, line_box):
            scores_in_block.append(score)
    return min(scores_in_block) if scores_in_block else 1.0

def add_paragraph_with_confidence_flag(doc, content, min_confidence, threshold=0.90):
    p = doc.add_paragraph()
    run = p.add_run(content)
    run.font.size = Pt(11)
    if min_confidence < threshold:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p

def sort_blocks_with_bbox_fallback(blocks):
    ordered = [b for b in blocks if b.order_index is not None]
    ordered.sort(key=lambda b: b.order_index)

    unordered = [b for b in blocks if b.order_index is None]

    result = list(ordered)
    for u_block in unordered:
        u_y = u_block.bbox[1]
        insert_pos = len(result)
        for idx, o_block in enumerate(result):
            if o_block.order_index is not None and o_block.bbox[1] > u_y:
                insert_pos = idx
                break
        result.insert(insert_pos, u_block)

    return result

def build_docx(image_path, output_path):
    pipeline = PPStructureV3(
        use_doc_orientation_classify=False,
        use_doc_unwarping=True
    )

    output = pipeline.predict(image_path)

    doc = Document()
    previous_page_index = None

    for res in output:
        current_page_index = res.get('page_index')

        # Insert a page break when moving to a new page (only for multi-page PDFs)
        if previous_page_index is not None and current_page_index != previous_page_index:
            doc.add_page_break()
        previous_page_index = current_page_index

        ocr_res = res['overall_ocr_res']
        ocr_lines = [
            (box, score) for box, text, score in zip(ocr_res['rec_boxes'], ocr_res['rec_texts'], ocr_res['rec_scores'])
            if text.strip()
        ]

        blocks = sort_blocks_with_bbox_fallback(res['parsing_res_list'])

        for block in blocks:
            label = block.label
            content = block.content.strip() if block.content else ""

            if not content:
                continue

            if label == 'paragraph_title':
                doc.add_heading(content, level=1)
            elif label == 'doc_title':
                doc.add_heading(content, level=0)
            elif label == 'table':
                add_html_table_to_doc(doc, content)
                doc.add_paragraph("")
            elif label == 'formula':
                render_formula_to_image(doc, content)
                doc.add_paragraph("")
            else:
                min_conf = get_min_confidence_for_block(block.bbox, ocr_lines)
                add_paragraph_with_confidence_flag(doc, content, min_conf)

    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == "__main__":
    build_docx('input_samples/test_multipage.pdf', 'output/multipage_test.docx')

if __name__ == "__main__":
    build_docx('input_samples/alice.png', 'output/confidence_test3.docx')