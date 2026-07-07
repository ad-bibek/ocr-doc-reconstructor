import streamlit as st
from paddleocr import PPStructureV3
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
import tempfile
import os
import zipfile
import io

st.set_page_config(page_title="OCR Doc Reconstructor")
st.title("Image to DOCX Converter")
st.write("Upload one or more document images and get back structured Word files.")

@st.cache_resource
def load_pipeline():
    return PPStructureV3(use_doc_orientation_classify=False, use_doc_unwarping=True)

def add_html_table_to_doc(doc, html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    rows = soup.find_all('tr')
    if not rows:
        return
    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < max_cols:
                table.cell(i, j).text = cell.get_text(strip=True)

def sort_blocks_with_bbox_fallback(blocks):
    ordered = [b for b in blocks if b.order_index is not None]
    ordered.sort(key=lambda b: b.order_index)

    unordered = [b for b in blocks if b.order_index is None]

    result = list(ordered)
    for u_block in unordered:
        u_y = u_block.bbox[1]
        insert_pos = len(result)
        for idx, o_block in enumerate(result):
            if o_block.order_index is not None and o_block.bbox[1] > u_y:
                insert_pos = idx
                break
        result.insert(insert_pos, u_block)

    return result

def add_page_to_doc(doc, output, is_first_page=True):
    """Adds one image's parsed content into an existing Document object."""
    for res in output:
        blocks = sort_blocks_with_bbox_fallback(res['parsing_res_list'])

        for block in blocks:
            label = block.label
            content = block.content.strip() if block.content else ""

            if not content:
                continue

            if label == 'paragraph_title':
                doc.add_heading(content, level=1)
            elif label == 'doc_title':
                doc.add_heading(content, level=0)
            elif label == 'table':
                add_html_table_to_doc(doc, content)
                doc.add_paragraph("")
            else:
                p = doc.add_paragraph(content)
                p.style.font.size = Pt(11)

def build_docx_single(image_path, output_path, pipeline):
    output = pipeline.predict(image_path)
    doc = Document()
    add_page_to_doc(doc, output)
    doc.save(output_path)

def build_docx_combined(image_paths, output_path, pipeline, progress_callback=None):
    doc = Document()
    for idx, image_path in enumerate(image_paths):
        output = pipeline.predict(image_path)
        if idx > 0:
            doc.add_page_break()
        add_page_to_doc(doc, output)
        if progress_callback:
            progress_callback(idx + 1, len(image_paths))
    doc.save(output_path)

uploaded_files = st.file_uploader(
    "Upload one or more images",
    type=['png', 'jpg', 'jpeg'],
    accept_multiple_files=True
)

if uploaded_files:
    st.write(f"**{len(uploaded_files)} image(s) uploaded**")

    cols = st.columns(min(len(uploaded_files), 4))
    for i, uf in enumerate(uploaded_files):
        with cols[i % 4]:
            st.image(uf, caption=uf.name, width='stretch')

    mode = st.radio(
        "Output mode",
        ["Separate DOCX files (zipped)", "Combined into one DOCX"],
        help="Separate: each image becomes its own Word file. Combined: all images become pages in one document, in upload order."
    )

    if st.button(f"Convert {len(uploaded_files)} image(s) to DOCX"):
        pipeline = load_pipeline()
        progress_bar = st.progress(0, text="Starting conversion...")

        with tempfile.TemporaryDirectory() as tmpdir:
            input_paths = []
            for uf in uploaded_files:
                p = os.path.join(tmpdir, uf.name)
                with open(p, 'wb') as f:
                    f.write(uf.getbuffer())
                input_paths.append(p)

            if mode == "Separate DOCX files (zipped)":
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w') as zf:
                    for i, (uf, input_path) in enumerate(zip(uploaded_files, input_paths)):
                        progress_bar.progress(
                            (i) / len(uploaded_files),
                            text=f"Processing {uf.name} ({i+1}/{len(uploaded_files)})..."
                        )
                        output_name = os.path.splitext(uf.name)[0] + ".docx"
                        output_path = os.path.join(tmpdir, output_name)
                        build_docx_single(input_path, output_path, pipeline)
                        zf.write(output_path, output_name)
                    progress_bar.progress(1.0, text="Done!")

                zip_buffer.seek(0)
                st.success("Done! All files converted.")
                st.download_button(
                    label="Download ZIP of converted DOCX files",
                    data=zip_buffer,
                    file_name="converted_documents.zip",
                    mime="application/zip"
                )

            else:
                output_path = os.path.join(tmpdir, "combined_output.docx")

                def update_progress(done, total):
                    progress_bar.progress(done / total, text=f"Processing image {done}/{total}...")

                build_docx_combined(input_paths, output_path, pipeline, update_progress)
                progress_bar.progress(1.0, text="Done!")

                with open(output_path, 'rb') as f:
                    docx_bytes = f.read()

                st.success("Done!")
                st.download_button(
                    label="Download combined DOCX",
                    data=docx_bytes,
                    file_name="combined_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )